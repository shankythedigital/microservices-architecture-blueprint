#!/usr/bin/env python3
"""Convert Markdown files to Word (.docx) with basic formatting."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def strip_md_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)


def clean_inline(text: str) -> str:
    text = strip_md_links(text)
    return text.replace("**", "").replace("`", "")


def add_formatted_run(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [clean_inline(cell.strip()) for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", line))


def set_document_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 5):
        name = f"Heading {level}"
        if name in doc.styles:
            h = doc.styles[name]
            h.font.name = "Calibri"
            h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9 if r_idx > 0 else 10)
                    if r_idx == 0:
                        r.bold = True
    doc.add_paragraph()


def convert_md_to_docx(md_path: Path, docx_path: Path, title: str | None = None) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    set_document_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        doc.add_paragraph()

    i = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if not is_table_separator(line):
                table_rows.append(parse_table_row(line))
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue

        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

        stripped = line.strip()

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(level, 4)
            heading_text = clean_inline(stripped.lstrip("#").strip())
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, clean_inline(re.sub(r"^[-*]\s+", "", stripped)))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, clean_inline(re.sub(r"^\d+\.\s+", "", stripped)))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_run(p, strip_md_links(line))
        i += 1

    if table_rows:
        add_table(doc, table_rows)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Created: {docx_path} ({docx_path.stat().st_size // 1024} KB)")


def main() -> None:
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py <input.md> <output.docx> [title]")
        sys.exit(1)
    md_path = Path(sys.argv[1]).resolve()
    docx_path = Path(sys.argv[2]).resolve()
    title = sys.argv[3] if len(sys.argv) > 3 else None
    convert_md_to_docx(md_path, docx_path, title)


if __name__ == "__main__":
    main()
